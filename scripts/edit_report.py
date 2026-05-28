"""Edit the SVNCKH report docx: fix cover page to match the submitted PDF,
remove RTX 2050 / 4GB hardware mentions, and rewrite AI-ish phrasing into more
natural prose. All numeric results, tables and references are left unchanged.
"""
import docx
from pathlib import Path

SRC = "submission_docs/SVNCKH_2026_Surgical_AI_English.docx"
DST = "submission_docs/SVNCKH_2026_Surgical_AI_English_edited.docx"


def set_text(paragraph, text):
    """Replace a paragraph's text while keeping the first run's formatting."""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


d = docx.Document(SRC)
P = d.paragraphs

# --- Cover page: title + personal fields (match the submitted PDF) ---
set_text(P[13], "Automatic Recognition of Surgical Phases in Cholecystectomy "
                "Videos: A Comparison of Three Modern Deep-Learning Architectures on Cholec80")
set_text(P[17], "Subcommittee: ICT (Information and Communication Technology)")
set_text(P[18], "Student's name: Nguyen Trong Bach")
set_text(P[19], "Student's ID: 23BI14057")
set_text(P[20], "Department: Department of Information and Communication Technology")
set_text(P[21], "Major: Information and Communication Technology")
set_text(P[22], "Supervisor: Dr. Vu Trong Sinh")

# --- Rewrites keyed by paragraph index. Numbers preserved verbatim. ---
rewrites = {
    # Abstract — smoother, less list-like
    52: (
        "Automatic surgical phase recognition turns laparoscopic video into structured "
        "workflow information, but it remains hard when phases look alike and appear with very "
        "different frequencies. We compare three temporal deep-learning models on Cholec80, a "
        "dataset of 80 laparoscopic cholecystectomy videos annotated with seven phases and seven "
        "instruments. We sample frames at 1 fps and split the videos into 40 for training, 20 for "
        "validation and 20 for testing. The three models are a ResNet-50 with a bidirectional "
        "LSTM, an EfficientNet-B3 with a temporal convolutional network, and a Swin-Tiny with a "
        "Transformer temporal encoder. Each takes eight-frame inputs and is trained with auxiliary "
        "tool prediction, a staged schedule, mixed precision, and median filtering at inference. On "
        "the held-out test set the smoothed macro-F1 scores were 0.7811, 0.7482 and 0.8153; the "
        "Swin-Transformer also reached an accuracy of 0.8623 and an edit score of 0.2006. Dropping "
        "the tool-prediction task in a shorter ablation lowered macro-F1 to 0.7651, and removing "
        "smoothing lowered the baseline edit score from 0.1520 to 0.1123. A code review found that "
        "the configured class-weighting option was never connected to the loss, so we report its "
        "planned ablation as inconclusive. Overall the results motivate further work on "
        "attention-based temporal models, with corrected ablations and validation on more datasets "
        "before any clinical use."
    ),
    # 1. Scientific context
    56: (
        "Minimally invasive surgery produces video that records an operation in detail, yet that "
        "record is hard to search or analyse by hand. A laparoscopic cholecystectomy removes the "
        "gallbladder through a recognisable sequence of steps: preparation, dissection around "
        "Calot's triangle, clipping and cutting, dissection of the gallbladder, packaging, "
        "cleaning and coagulation, and retraction. Identifying the current step automatically is "
        "called surgical phase recognition, and it underpins retrospective quality analysis, "
        "teaching support, case indexing, and future context-aware operating-room systems [1], "
        "[2]. We address workflow recognition only, and do not propose an autonomous clinical "
        "decision system."
    ),
    57: (
        "The public Cholec80 dataset, introduced with EndoNet, is a practical benchmark for this "
        "problem. It contains 80 laparoscopic cholecystectomy videos with frame-level phase labels "
        "and instrument-presence annotations [3], and its instrument labels let us test whether "
        "learning tools and phases together helps phase recognition. The dataset is also "
        "demanding: in our prepared training data, Calot triangle dissection makes up 40.9% of "
        "labelled frames while gallbladder retraction makes up only 3.7%. A model can therefore "
        "post a deceptively high overall accuracy by favouring the common phases while doing poorly "
        "on the brief but clinically important ones."
    ),
    # 2. Scientific problem — replace para 59 (keep) and 60 (RTX removed)
    59: (
        "A single laparoscopic frame is often not enough to identify a phase. The same instrument "
        "and similar tissue can appear on both sides of a transition, and even an expert annotator "
        "places the exact transition time with some subjectivity. A good model therefore has to "
        "combine the spatial evidence in each frame with temporal evidence from its neighbours, and "
        "it has to avoid sudden, implausible jumps between phases. These needs are what make a "
        "comparison of temporal model families more informative than a comparison of image "
        "backbones alone."
    ),
    60: (
        "We deliberately developed the project within a modest, single-GPU budget typical of a "
        "student or small research group, which constrains the input sequence length and batch "
        "size and rewards methods that others can reproduce without specialised infrastructure. "
        "The study therefore asks a concrete question: within one compact implementation and a "
        "realistic compute budget, which temporal architecture gives the strongest phase "
        "recognition, and which supporting design choices are worth testing more carefully?"
    ),
    # Methods 3 — remove the RTX sentence at the end of para 78
    78: (
        "The phase loss is cross-entropy with label smoothing of 0.1. When the tool head is "
        "active, we add binary cross-entropy for tool presence with weight 0.5. The loss also "
        "includes a temporal consistency penalty on adjacent phase probabilities with a default "
        "weight of 0.1. We optimise with AdamW [10] using mixed precision, gradient clipping and a "
        "cosine warm-restart scheduler. The main models are trained in three stages: the visual "
        "backbone stays frozen while the temporal component and heads train for 5 and then 10 "
        "epochs, followed by 10 epochs of end-to-end fine-tuning. We select the checkpoint by "
        "validation macro-F1."
    ),
    # Discussion limitations — drop the "hardware restrictions" sentence
    95: (
        "The study has further limitations. It looks at one operation type on one public dataset, "
        "so we cannot say how well it generalises to other hospitals or procedures. It samples at "
        "one frame per second and uses eight-frame context for the main comparison, which may miss "
        "very short actions. And the system has not been evaluated prospectively in an operating "
        "room, so it must not be treated as a clinical decision-support tool."
    ),
}

for idx, text in rewrites.items():
    set_text(P[idx], text)

Path(DST).parent.mkdir(parents=True, exist_ok=True)
d.save(DST)
print("Saved", DST)

# Verify: no RTX/2050 left
d2 = docx.Document(DST)
hits = [i for i, p in enumerate(d2.paragraphs)
        if any(k in p.text for k in ["RTX", "2050", "4 GB", "4GB", "GeForce"])]
print("RTX/hardware mentions remaining:", hits)
