"""Fill in the supervisor-comment block in the edited report docx with a version
that highlights the practical application of the work, keeping the same 4-paragraph
length and examiner voice. Also build a Vietnamese reading copy of the comment.
"""
import docx

DOC = "submission_docs/SVNCKH_2026_Surgical_AI_English_edited.docx"

COMMENT_EN = [
    "The report presents a clear and relevant research topic in the field of surgical "
    "video analysis and deep learning. The student demonstrates a good understanding of "
    "the scientific problem, especially the difficulty of surgical phase recognition due "
    "to class imbalance, visually similar phases, and the need for temporal modelling.",

    "The work shows a serious effort to build and compare different deep-learning "
    "architectures, including recurrent, temporal convolutional, and Transformer-based "
    "models. Beyond the comparison itself, the topic has clear practical value: reliable "
    "phase recognition can support surgical training and skill assessment, automatic "
    "indexing and retrieval of operative video, retrospective quality and workflow "
    "analysis, and, in the longer term, context-aware assistance in the operating room. "
    "The experimental design is appropriate for a student research project, with a clear "
    "dataset split and meaningful evaluation metrics.",

    "The report also shows scientific honesty by identifying implementation limitations "
    "and avoiding unsupported conclusions. The results are promising and provide a useful "
    "basis for further development and real deployment, especially in improving temporal "
    "inference, completing ablation studies, and validating the method on additional "
    "surgical datasets and hospitals before any clinical use.",

    "Overall, this is a well-prepared and technically sound report. The work demonstrates "
    "initiative, careful implementation, and good potential for continuation as a student "
    "scientific research project with a concrete path toward applied surgical-workflow tools.",
]

# Vietnamese reading copy (not inserted into the doc; saved as a separate file)
COMMENT_VI = [
    "Báo cáo trình bày một đề tài nghiên cứu rõ ràng và phù hợp trong lĩnh vực phân tích "
    "video phẫu thuật và học sâu. Sinh viên thể hiện sự hiểu biết tốt về bài toán khoa "
    "học, đặc biệt là khó khăn của nhận dạng giai đoạn phẫu thuật do mất cân bằng lớp, các "
    "giai đoạn có hình ảnh tương tự nhau, và yêu cầu mô hình hoá theo thời gian.",

    "Công trình cho thấy nỗ lực nghiêm túc trong việc xây dựng và so sánh các kiến trúc "
    "học sâu khác nhau, bao gồm mô hình hồi quy (recurrent), tích chập theo thời gian "
    "(temporal convolutional) và mô hình dựa trên Transformer. Ngoài bản thân việc so sánh, "
    "đề tài còn có giá trị ứng dụng rõ ràng: nhận dạng giai đoạn đáng tin cậy có thể hỗ trợ "
    "đào tạo và đánh giá kỹ năng phẫu thuật, tự động lập chỉ mục và truy xuất video ca mổ, "
    "phân tích chất lượng và quy trình sau mổ, và về lâu dài là hỗ trợ theo ngữ cảnh ngay "
    "trong phòng mổ. Thiết kế thí nghiệm phù hợp với một đề tài nghiên cứu sinh viên, với "
    "cách chia tập dữ liệu rõ ràng và các chỉ số đánh giá có ý nghĩa.",

    "Báo cáo cũng thể hiện tính trung thực khoa học khi chỉ ra những hạn chế trong cài đặt "
    "và tránh đưa ra các kết luận thiếu căn cứ. Kết quả khả quan và là một nền tảng hữu ích "
    "cho việc phát triển tiếp theo cũng như triển khai thực tế, đặc biệt ở các hướng cải "
    "thiện suy luận theo thời gian, hoàn thiện các thí nghiệm loại bỏ (ablation), và kiểm "
    "chứng phương pháp trên các tập dữ liệu phẫu thuật và bệnh viện khác trước khi ứng dụng "
    "lâm sàng.",

    "Nhìn chung, đây là một báo cáo được chuẩn bị tốt và vững về mặt kỹ thuật. Công trình "
    "thể hiện tính chủ động, sự cẩn thận trong cài đặt, và tiềm năng tốt để tiếp tục như "
    "một đề tài nghiên cứu khoa học sinh viên, với một lộ trình cụ thể hướng tới các công "
    "cụ phân tích quy trình phẫu thuật ứng dụng được trong thực tế.",
]

d = docx.Document(DOC)
P = d.paragraphs

# The comment block occupies the empty paragraphs between index 31 and 47.
# Write the 4 paragraphs into the first 4 empty slots (32, 34, 36, 38) so they
# keep visual spacing; clear any stray text in the others up to 46.
slots = [32, 34, 36, 38]
for slot, text in zip(slots, COMMENT_EN):
    p = P[slot]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)

d.save(DOC)
print("Updated supervisor comment in", DOC)

# Write a standalone Vietnamese reading copy
with open("submission_docs/Nhan_xet_giang_vien_TiengViet.txt", "w", encoding="utf-8") as f:
    f.write("NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN VỀ ĐÓNG GÓP KHOA HỌC CỦA BÁO CÁO\n")
    f.write("(Phần này do giảng viên hướng dẫn ghi)\n\n")
    for para in COMMENT_VI:
        f.write(para + "\n\n")
    f.write("Hà Nội, ngày 28 tháng 5 năm 2026\n")
    f.write("Giảng viên hướng dẫn\n(Ký và ghi rõ họ tên)\n\nVũ Trọng Sinh\n")
print("Wrote Vietnamese copy")
